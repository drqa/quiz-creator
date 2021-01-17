from Question import Question
import random
import PySimpleGUI as sg
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

q_num = 1
abc = ["        A. ", "        B. ", "        C. ", "        D. "]
questions = []
keys_to_clear = ['question', 'a', 'b', 'c', 'd', 'confirm']

sg.theme("Reddit")
layout = [[sg.Text("Save As: "), sg.InputText(size=(30,1),key='name')],
          [sg.Text("#" + str(q_num), key='number',size=(3,1)), sg.Multiline(size=(60,4), key='question')],
          [sg.Text("A.",size=(3,1)), sg.Multiline(size=(60,2), key='a')],
          [sg.Text("B.",size=(3,1)), sg.Multiline(size=(60,2), key='b')],
          [sg.Text("C.",size=(3,1)), sg.Multiline(size=(60,2), key='c')],
          [sg.Text("D.",size=(3,1)), sg.Multiline(size=(60,2), key='d')],
          [sg.Button("Add Question")], 
          [sg.Button("Create"), sg.InputText(key='versions', size=(3,1)), sg.Text("versions"), sg.Text("",size=(35,1),key='confirm', text_color="blue")],
          [sg.Button("Close")]]

window = sg.Window("Shuffler", layout, finalize=True)

while True:
    event, values = window.read()
    if event == "Add Question":
        questions.append(Question(values['question'],[values['a'],values['b'],values['c'],values['d']]))
        for key in keys_to_clear:
            window[key]('')
        q_num+=1
        window['number']("#" + str(q_num))

    if event == "Create":
        name = values['name']
        i = 0
        j = 0
        x = 0
        while x < int(values['versions']):
            quiz = docx.Document()
            title = quiz.add_paragraph(name + " Version: " + str(x + 1))
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_format = quiz.styles['Normal'].paragraph_format
            paragraph_format.line_spacing = 1
            random.shuffle(questions)

            for question in questions:
                random.shuffle(question.answers)
                question_para = quiz.add_paragraph(str(i + 1) + "." + " " + str(question.body))

                while j < 4:
                    answer_para = quiz.add_paragraph(abc[j]).add_run(question.answers[j])
                    j += 1
                j = 0
                i += 1
            x += 1
            i = 0
            quiz.save('/home/daniel/Documents/' + name + str(x) + '.docx')  
            window['confirm']('File Creation Successful')
    if event == sg.WIN_CLOSED or event == "Close":
        break

window.close()
